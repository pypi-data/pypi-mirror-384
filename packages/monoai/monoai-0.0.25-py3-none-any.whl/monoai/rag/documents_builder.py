import os
import uuid
import importlib
import importlib.util
from typing import List, Dict, Tuple, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document as _DocxDocument
    import docx2txt as _docx2txt
    import PyPDF2 as _PyPDF2


class DocumentsBuilder:
    """
    A utility class for building document collections from various sources.
    
    This class provides methods to extract text content from files and web pages,
    split the content into manageable chunks with configurable size and overlap,
    and prepare the data for storage in vector databases.
    
    The DocumentsBuilder is designed to work seamlessly with the RAG system,
    producing output that can be directly used with vector database operations.
    
    Features:
    - File-based document extraction with UTF-8 encoding support
    - Text string processing for in-memory content
    - Web scraping with multiple engine options (requests, tavily, selenium)
    - Word document extraction (.doc and .docx formats)
    - PDF document extraction with metadata
    - Multiple chunking strategies (word, sentence, paragraph, fixed, semantic)
    - Configurable chunk size and overlap parameters
    - Rich metadata generation for each document chunk
    - Unique ID generation for database storage
    
    Attributes:
        _chunk_strategy (str): The chunking strategy to use
        _chunk_size (int): Maximum size of each text chunk in characters
        _chunk_overlap (int): Number of characters to overlap between chunks
    """

    def __init__(
        self,
        chunk_strategy: str = "word",
        chunk_size: int = 1000,
        chunk_overlap: int = 0,
        custom_split_func: Optional[callable] = None
    ):
        """
        Initialize the DocumentsBuilder with chunking parameters.
        """
        # If custom_split_func is provided, automatically set strategy to "custom"
        if custom_split_func is not None:
            chunk_strategy = "custom"
        
        self._chunk_strategy = chunk_strategy
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._custom_split_func = custom_split_func

        self._module_cache: Dict[str, object] = {}
        
        if chunk_overlap >= chunk_size:
            raise ValueError(
                f"chunk_overlap ({chunk_overlap}) must be less than chunk_size ({chunk_size}) "
                "to prevent infinite loops. Recommended: chunk_overlap should be 10-20% of chunk_size."
            )
        
        if chunk_size <= 0:
            raise ValueError(f"chunk_size must be positive, got {chunk_size}")
        
        if chunk_overlap < 0:
            raise ValueError(f"chunk_overlap must be non-negative, got {chunk_overlap}")
        
        if chunk_strategy == "custom" and custom_split_func is None:
            raise ValueError("custom_split_func must be provided when chunk_strategy='custom'")
        
        if custom_split_func is not None and not callable(custom_split_func):
            raise ValueError("custom_split_func must be callable")

    # ---------- Lazy import helpers ----------
    @staticmethod
    def _has_module(name: str) -> bool:
        """Rileva se un modulo è installato senza importarlo completamente."""
        spec = importlib.util.find_spec(name)
        return spec is not None

    def _get_module(self, name: str):
        """Importa pigramente un modulo e lo cache-izza."""
        mod = self._module_cache.get(name)
        if mod is None:
            try:
                mod = importlib.import_module(name)
            except ImportError as e:
                raise ImportError(
                    f"Il modulo '{name}' è richiesto per questa operazione. "
                    f"Installa con: pip install {name}"
                ) from e
            self._module_cache[name] = mod
        return mod

    def from_file(self, file_path: str) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Read a file and split it into chunks with specified size and overlap.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        chunks = self._split_text(text)
        
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_str(self, text: str, source_name: str = "text_string") -> Tuple[List[str], List[Dict], List[str]]:
        """
        Process a text string and split it into chunks with specified size and overlap.
        """
        if not text or not text.strip():
            return [], [], []
        
        chunks = self._split_text(text)
        
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            metadata = {
                'source_type': 'text_string',
                'source_name': source_name,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk),
                'chunk_strategy': self._chunk_strategy
            }
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_doc(self, file_path: str, extraction_method: str = "auto") -> Tuple[List[str], List[Dict], List[str]]:
        """
        Extract text from Word documents (.doc and .docx files) and split into chunks.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in ['.doc', '.docx']:
            raise ValueError(f"Unsupported file format: {file_extension}. Only .doc and .docx files are supported.")
        
        # Selezione metodo al volo con lazy detection
        if extraction_method == "auto":
            if file_extension == '.docx' and self._has_module("docx"):
                extraction_method = "docx"
            elif self._has_module("docx2txt"):
                extraction_method = "docx2txt"
            else:
                raise ImportError(
                    "docx2txt is required for .docx file extraction. Install with: pip install docx2txt"
                )
        
        if extraction_method == "docx":
            if file_extension != '.docx':
                raise ValueError("'docx' extraction method only supports .docx files")
            text, doc_properties = self._extract_with_docx(file_path)
        elif extraction_method == "docx2txt":
            text, doc_properties = self._extract_with_docx2txt(file_path)
        else:
            raise ValueError(f"Unsupported extraction method: {extraction_method}")
        
        chunks = self._split_text(text)
        
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'document_format': file_extension[1:],  # Remove the dot
                'extraction_method': extraction_method,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            if doc_properties:
                metadata.update(doc_properties)
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_pdf(self, file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Extract text from PDF documents and split into chunks.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension != '.pdf':
            raise ValueError(f"Unsupported file format: {file_extension}. Only .pdf files are supported.")
        
        # qui PyPDF2 viene importato solo se davvero richiesto
        if not self._has_module("PyPDF2"):
            raise ImportError("PyPDF2 is required for PDF file extraction. Install with: pip install PyPDF2")
        
        text, pdf_properties, page_info = self._extract_from_pdf(file_path, page_range)
        
        chunks = self._split_text(text)
        
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'document_format': 'pdf',
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            if pdf_properties:
                metadata.update(pdf_properties)
            if page_info:
                metadata.update(page_info)
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_url(self, url: str, engine: str = "requests", deep: bool = False) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Scrape content from a URL and split it into chunks with specified size and overlap.
        
        Notes:
        - Scraping may take time depending on the engine and website complexity
        - Some websites may block automated scraping
        - Selenium requires Chrome/Chromium to be installed
        - Tavily requires an API key to be configured
        """
        from monoai.tools.webscraping import scrape_web
        result = scrape_web(url, engine=engine, deep=deep)
                
        if not result or not result.get("text"):
            raise ValueError(f"Failed to extract text content from URL: {url}")
        
        text = result["text"]
        chunks = self._split_text(text)
        
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            metadata = {
                'url': url,
                'source_type': 'web_page',
                'scraping_engine': engine,
                'deep_extraction': deep,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids
    
    def _extract_with_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from a .docx file using python-docx library (lazy import).
        """
        docx = self._get_module("docx")
        Document = docx.Document 
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n\n".join(text_parts)
        
        # Extract document properties
        properties = {}
        core_props = doc.core_properties
        if core_props.title:
            properties['document_title'] = core_props.title
        if core_props.author:
            properties['document_author'] = core_props.author
        if core_props.subject:
            properties['document_subject'] = core_props.subject
        if core_props.created:
            properties['document_created'] = str(core_props.created)
        if core_props.modified:
            properties['document_modified'] = str(core_props.modified)
        
        return text, properties
    
    def _extract_with_docx2txt(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from a Word document using docx2txt library (lazy import).
        """
        docx2txt = self._get_module("docx2txt")
        text = docx2txt.process(file_path)  
        return text, {}
    
    def _extract_from_pdf(self, file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[str, Dict, Dict]:
        """
        Extract text and metadata from a PDF file using PyPDF2 (lazy import).
        """
        PyPDF2 = self._get_module("PyPDF2")
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file) 
            
            # Get total number of pages
            total_pages = len(pdf_reader.pages)
            
            # Determine page range
            if page_range is None:
                start_page = 1
                end_page = total_pages
            else:
                start_page, end_page = page_range
                # Validate page range
                if start_page < 1 or end_page > total_pages or start_page > end_page:
                    raise ValueError(f"Invalid page range: {page_range}. Pages must be between 1 and {total_pages}")
            
            # Extract text from specified pages
            text_parts = []
            for page_num in range(start_page - 1, end_page):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            # Extract PDF properties
            properties = {}
            if getattr(pdf_reader, "metadata", None):
                metadata = pdf_reader.metadata

                def _get(meta, key):
                    try:
                        return meta.get(key) if hasattr(meta, "get") else getattr(meta, key, None)
                    except Exception:
                        return None

                mapping = [
                    ('/Title', 'pdf_title'),
                    ('/Author', 'pdf_author'),
                    ('/Subject', 'pdf_subject'),
                    ('/Creator', 'pdf_creator'),
                    ('/Producer', 'pdf_producer'),
                    ('/CreationDate', 'pdf_creation_date'),
                    ('/ModDate', 'pdf_modification_date'),
                ]
                for k_src, k_dst in mapping:
                    val = _get(metadata, k_src)
                    if val:
                        properties[k_dst] = str(val)
            
            # Add page information
            page_info = {
                'total_pages': total_pages,
                'extracted_pages_start': start_page,
                'extracted_pages_end': end_page,
                'extracted_pages_count': end_page - start_page + 1
            }
            
            return text, properties, page_info
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks using the specified chunking strategy.
        """
        if len(text) <= self._chunk_size:
            return [text]
        
        if self._chunk_strategy == "word":
            return self._split_by_words(text)
        elif self._chunk_strategy == "sentence":
            return self._split_by_sentences(text)
        elif self._chunk_strategy == "paragraph":
            return self._split_by_paragraphs(text)
        elif self._chunk_strategy == "fixed":
            return self._split_fixed(text)
        elif self._chunk_strategy == "semantic":
            return self._split_semantic(text)
        elif self._chunk_strategy == "custom":
            return self._custom_split_func(text, self._chunk_size, self._chunk_overlap)
        else:
            raise ValueError(f"Unsupported chunk strategy: {self._chunk_strategy}")
    
    def _split_by_words(self, text: str) -> List[str]:
        """
        Split text by word boundaries while respecting word count.
        """
        words = text.split()
        
        if len(words) <= self._chunk_size:
            return [text]
        
        chunks = []
        start_word = 0
        
        while start_word < len(words):
            end_word = start_word + self._chunk_size
            chunk_words = words[start_word:end_word]
            chunk = ' '.join(chunk_words)
            
            if chunk.strip():
                chunks.append(chunk)
            
            new_start_word = end_word - self._chunk_overlap
            if new_start_word <= start_word:
                new_start_word = start_word + 1
            start_word = new_start_word
            
            if start_word >= len(words):
                break
        
        return chunks
    
    def _split_by_sentences(self, text: str) -> List[str]:
        """
        Split text by sentence boundaries while respecting sentence count.
        """
        sentence_endings = ['.', '!', '?', '\n\n']
        
        sentences = []
        last_pos = 0
        
        for i, char in enumerate(text):
            if char in sentence_endings:
                sentence = text[last_pos:i+1].strip()
                if sentence:
                    sentences.append(sentence)
                last_pos = i + 1
        
        if last_pos < len(text):
            last_sentence = text[last_pos:].strip()
            if last_sentence:
                sentences.append(last_sentence)
        
        if len(sentences) <= self._chunk_size:
            return [text]
        
        chunks = []
        start_sentence = 0
        
        while start_sentence < len(sentences):
            end_sentence = start_sentence + self._chunk_size
            chunk_sentences = sentences[start_sentence:end_sentence]
            chunk = ' '.join(chunk_sentences)
            
            if chunk.strip():
                chunks.append(chunk)
            
            new_start_sentence = end_sentence - self._chunk_overlap
            if new_start_sentence <= start_sentence:
                new_start_sentence = start_sentence + 1
            start_sentence = new_start_sentence
            
            if start_sentence >= len(sentences):
                break
        
        return chunks
    
    def _split_by_paragraphs(self, text: str) -> List[str]:
        """
        Split text by paragraph boundaries while respecting paragraph count.
        """
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if len(paragraphs) <= self._chunk_size:
            return [text]
        
        chunks = []
        start_paragraph = 0
        
        while start_paragraph < len(paragraphs):
            end_paragraph = start_paragraph + self._chunk_size
            chunk_paragraphs = paragraphs[start_paragraph:end_paragraph]
            chunk = '\n\n'.join(chunk_paragraphs)
            
            if chunk.strip():
                chunks.append(chunk)
            
            new_start_paragraph = end_paragraph - self._chunk_overlap
            if new_start_paragraph <= start_paragraph:
                new_start_paragraph = start_paragraph + 1
            start_paragraph = new_start_paragraph
            
            if start_paragraph >= len(paragraphs):
                break
        
        return chunks
    
    def _split_fixed(self, text: str) -> List[str]:
        """
        Split text into fixed-size chunks without considering boundaries.
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self._chunk_size
            chunk = text[start:end].strip()
            
            if chunk:
                chunks.append(chunk)
            
            new_start = end - self._chunk_overlap
            if new_start <= start:
                new_start = start + 1
            start = new_start
            
            if start >= len(text):
                break
        
        return chunks
    
    def _split_semantic(self, text: str) -> List[str]:
        """
        Split text by semantic boundaries.
        """
        semantic_patterns = [
            '\n# ', '\n## ', '\n### ', '\n#### ',  # Markdown headers
            '\n1. ', '\n2. ', '\n3. ', '\n4. ', '\n5. ',  # Numbered lists
            '\n• ', '\n- ', '\n* ',  # Bullet points
            '\n\n',  # Paragraph breaks
            '\n---\n', '\n___\n',  # Horizontal rules
            '\n\nChapter ', '\n\nSection ', '\n\nPart ',  # Document sections
        ]
        
        chunks = []
        current_chunk = ""
        
        parts = [text]
        for pattern in semantic_patterns:
            new_parts = []
            for part in parts:
                if pattern in part:
                    split_parts = part.split(pattern)
                    for i, split_part in enumerate(split_parts):
                        if i > 0:
                            split_part = pattern + split_part
                        if split_part.strip():
                            new_parts.append(split_part)
                else:
                    new_parts.append(part)
            parts = new_parts
        
        for part in parts:
            if len(current_chunk) + len(part) > self._chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                overlap_start = max(0, len(current_chunk) - self._chunk_overlap)
                current_chunk = current_chunk[overlap_start:] + part
            else:
                current_chunk += part
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
