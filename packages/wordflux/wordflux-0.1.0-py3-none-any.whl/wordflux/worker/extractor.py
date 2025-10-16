from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from dataclasses import asdict
import json
import zipfile
from lxml import etree
from wordflux.document.document import TextSegment, TableCellSegment, ChartSegment, SmartArtSegment, RunInfo
from wordflux.utils.is_numeric import is_numeric
from wordflux.utils.decorator import progress_tracker, timer, log_errors
import logging

logging.basicConfig(level=logging.WARNING)

class Extractor:
    """Tr√≠ch xu·∫•t vƒÉn b·∫£n t·ª´ ƒëo·∫°n vƒÉn, √¥ b·∫£ng, bi·ªÉu ƒë·ªì v√† SmartArt t·ª´ file DOCX"""
    
    def __init__(self, input_file: str, checkpoint_file: str):
        self.input_file = input_file
        self.checkpoint_file = checkpoint_file
        self.doc = Document(input_file)
        self.text_segments = []
        self.table_cell_segments = []
        self.chart_segments = []
        self.smartart_segments = []
        self.logger = logging.getLogger(self.__class__.__name__)
        
        self.ns = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
        }

    def _has_smartart_or_chart(self, para: Paragraph):
        """Ki·ªÉm tra xem ƒëo·∫°n vƒÉn c√≥ ch·ª©a SmartArt ho·∫∑c Chart kh√¥ng"""
        para_elem = para._element
        for run_elem in para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            drawing = run_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawing is not None:
                return True
        return False
    
    def _extract_runs(self, para: Paragraph):
        """Tr√≠ch xu·∫•t v√† g·ªôp c√°c run c√≥ c√πng ƒë·ªãnh d·∫°ng"""
        runs_list = []
        for run in para.runs:
            if run.text == "":
                continue
            run_info = RunInfo(
                text=run.text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                superscript=run.font.superscript,
                subscript=run.font.subscript,
            )
            # G·ªôp c√°c run li√™n ti·∫øp c√≥ c√πng ƒë·ªãnh d·∫°ng
            if runs_list and run_info == runs_list[-1]:
                runs_list[-1].text += run.text
            else:
                runs_list.append(run_info)
        return runs_list
    
    # @progress_tracker(item_name='paragraphs', use_tqdm=True)
    def _extract_text_segments(self, paragraphs: list[Paragraph], progress_callback=None):
        """Tr√≠ch xu·∫•t t·∫•t c·∫£ c√°c ƒëo·∫°n vƒÉn b·∫£n th√¥ng th∆∞·ªùng"""
        for seg_idx, para in enumerate(paragraphs):
            full_text = para.text.strip()
            if full_text:
                text_segment = TextSegment(seg_idx, full_text, self._has_smartart_or_chart(para))
                text_segment.runs_list = self._extract_runs(para)
                self.text_segments.append(text_segment)
            
            if progress_callback:
                progress_callback()
    
    # @progress_tracker(item_name='table cells', use_tqdm=True)
    def _extract_table_cell_segments(self, tables: list[Table], progress_callback=None):
        """Tr√≠ch xu·∫•t vƒÉn b·∫£n t·ª´ t·∫•t c·∫£ c√°c √¥ trong b·∫£ng"""
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        full_text = para.text.strip()
                        if full_text:
                            table_cell_segment = TableCellSegment(table_idx, row_idx, cell_idx, para_idx)
                            table_cell_segment.runs_list = self._extract_runs(para)
                            self.table_cell_segments.append(table_cell_segment)
                        
                        if progress_callback:
                            progress_callback()
    
    # @progress_tracker(item_name='chart files', use_tqdm=True)
    def _extract_chart_segments(self, chart_files: list[str], progress_callback=None):
        """Tr√≠ch xu·∫•t vƒÉn b·∫£n t·ª´ bi·ªÉu ƒë·ªì (ti√™u ƒë·ªÅ, gi√° tr·ªã, danh m·ª•c)"""
        try:
            with zipfile.ZipFile(self.input_file) as z:
                for chart_idx, chart_file in enumerate(chart_files):
                    try:
                        chart_xml = z.read(chart_file)
                        chart_root = etree.fromstring(chart_xml)
                        
                        # Tr√≠ch xu·∫•t ti√™u ƒë·ªÅ bi·ªÉu ƒë·ªì
                        titles = chart_root.findall('.//c:title', self.ns)
                        for title_idx, title in enumerate(titles):
                            text_elems = title.findall('.//a:t', self.ns)
                            for t_elem in text_elems:
                                if t_elem.text and t_elem.text.strip():
                                    self.chart_segments.append(ChartSegment(
                                        chart_idx=chart_idx,
                                        element_type="title",
                                        element_idx=title_idx,
                                        text=t_elem.text.strip(),
                                        file_path=chart_file
                                    ))
                        
                        # Tr√≠ch xu·∫•t gi√° tr·ªã trong bi·ªÉu ƒë·ªì
                        v_elements = chart_root.findall('.//c:v', self.ns)
                        v_idx = 0
                        for v_elem in v_elements:
                            if v_elem.text and v_elem.text.strip():
                                text = v_elem.text.strip()
                                if not is_numeric(text):
                                    self.chart_segments.append(ChartSegment(
                                        chart_idx=chart_idx,
                                        element_type="value",
                                        element_idx=v_idx,
                                        text=text,
                                        file_path=chart_file
                                    ))
                                    v_idx += 1
                    except Exception as e:
                        self.logger.warning(f"   ‚ö†Ô∏è  Error processing {chart_file}: {e}")
                    finally:
                        if progress_callback:
                            progress_callback()
        except Exception as e:
            self.logger.warning(f"   ‚ö†Ô∏è  Error accessing charts: {e}")
    
    # @progress_tracker(item_name='SmartArt files', use_tqdm=True)
    def _extract_smartart_segments(self, diagram_files: list[str], progress_callback=None):
        """Tr√≠ch xu·∫•t vƒÉn b·∫£n t·ª´ SmartArt"""
        try:
            with zipfile.ZipFile(self.input_file) as z:
                for smartart_idx, diagram_file in enumerate(diagram_files):
                    try:
                        diagram_xml = z.read(diagram_file)
                        diagram_root = etree.fromstring(diagram_xml)
                        text_elems = diagram_root.findall('.//a:t', self.ns)
                        
                        for elem_idx, text_elem in enumerate(text_elems):
                            if text_elem.text and text_elem.text.strip():
                                self.smartart_segments.append(SmartArtSegment(
                                    smartart_idx=smartart_idx,
                                    element_idx=elem_idx,
                                    text=text_elem.text.strip(),
                                    file_path=diagram_file
                                ))
                    except Exception as e:
                        self.logger.warning(f"   ‚ö†Ô∏è  Error processing {diagram_file}: {e}")
                    finally:
                        if progress_callback:
                            progress_callback()
        except Exception as e:
            self.logger.warning(f"   ‚ö†Ô∏è  Error accessing SmartArt: {e}")
    
    @timer
    @log_errors
    def extract(self):
        """Th·ª±c hi·ªán tr√≠ch xu·∫•t to√†n b·ªô n·ªôi dung v√† l∆∞u checkpoint"""
        self.logger.info("="*70)
        self.logger.info("EXTRACTING ALL CONTENT FROM DOCX")
        self.logger.info("="*70 + "\n")
        
        # Extract v·ªõi progress bars
        self._extract_text_segments(self.doc.paragraphs)
        
        # ƒê·∫øm t·ªïng s·ªë cells trong t·∫•t c·∫£ c√°c b·∫£ng
        all_cells = []
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_cells.append((table, row, cell, para))
        self._extract_table_cell_segments(self.doc.tables)
        
        # L·∫•y danh s√°ch chart v√† SmartArt files
        try:
            with zipfile.ZipFile(self.input_file) as z:
                chart_files = [f for f in z.namelist() if 'chart' in f.lower() and f.endswith('.xml')]
                diagram_files = [f for f in z.namelist() if 'diagram' in f.lower() and f.endswith('.xml')]
                
                if chart_files:
                    self._extract_chart_segments(chart_files)
                else:
                    self.logger.info("üìà No chart files found")
                
                if diagram_files:
                    self._extract_smartart_segments(diagram_files)
                else:
                    self.logger.info("üï∏Ô∏è No SmartArt files found")
        except Exception as e:
            self.logger.warning(f"‚ö†Ô∏è  Error accessing ZIP content: {e}")
        
        # L∆∞u checkpoint
        with open(self.checkpoint_file, "w", encoding="utf-8") as f:
            data = {
                "text_segments": [asdict(segment) for segment in self.text_segments],
                "table_cell_segments": [asdict(segment) for segment in self.table_cell_segments],
                "chart_segments": [asdict(segment) for segment in self.chart_segments],
                "smartart_segments": [asdict(segment) for segment in self.smartart_segments],
            }
            json.dump(data, f, ensure_ascii=False, indent=2)