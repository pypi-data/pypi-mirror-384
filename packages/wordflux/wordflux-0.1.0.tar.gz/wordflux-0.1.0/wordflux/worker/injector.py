import json
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx import Document
from wordflux.document.document import RunInfo, TextSegment, TableCellSegment, ChartSegment, SmartArtSegment
import os
import zipfile
import shutil
from lxml import etree
import tempfile
from wordflux.utils.is_numeric import is_numeric
from wordflux.utils.decorator import progress_tracker, timer, log_errors
import logging

logging.basicConfig(level=logging.WARNING)

class Injector:
    """Chèn nội dung đã dịch trở lại vào file DOCX"""
    
    def __init__(self, input_file: str, checkpoint_file: str, output_file: str):
        self.input_file = input_file
        self.checkpoint_file = checkpoint_file
        self.output_file = output_file
        self.doc = Document(input_file)
        self.logger = logging.getLogger(self.__class__.__name__)
        
        self.ns = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
        }
    
    def _apply_runs(self, para: Paragraph, runs_list: list[RunInfo]):
        """Áp dụng danh sách runs vào paragraph"""
        for run_info in runs_list:
            new_run = para.add_run(run_info['translated_text'])
            new_run.bold = run_info['bold']
            new_run.italic = run_info['italic']
            new_run.underline = run_info['underline']
            new_run.font.superscript = run_info['superscript']
            new_run.font.subscript = run_info['subscript']
    
    def _clear_except_important(self, para):
        """Xóa văn bản trong đoạn văn nhưng giữ lại SmartArt/Chart (các run chứa <w:drawing>)"""
        para_elem = para._element
        runs_to_keep = []
        for run_elem in para_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            if run_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                runs_to_keep.append(run_elem)
        
        para.clear()
        for run_elem in runs_to_keep:
            para_elem.append(run_elem)
    
    # @progress_tracker(item_name='text segments', use_tqdm=True)
    def _inject_text_segments(self, text_segments: list[TextSegment], progress_callback=None):
        """Chèn các đoạn văn bản"""
        for text_segment in text_segments:
            para = self.doc.paragraphs[text_segment["seg_idx"]]
            if text_segment["has_smartart_or_chart"]:
                self._clear_except_important(para)
            else:
                para.clear()
            self._apply_runs(para, text_segment["runs_list"])
            
            if progress_callback:
                progress_callback()
    
    # @progress_tracker(item_name='table cells', use_tqdm=True)
    def _inject_table_cell_segments(self, table_cell_segments: list[TableCellSegment], progress_callback=None):
        """Chèn văn bản vào ô bảng"""
        for segment in table_cell_segments:
            table = self.doc.tables[segment['table_idx']]
            row = table.rows[segment['row_idx']]
            cell = row.cells[segment['cell_idx']]
            para = cell.paragraphs[segment['para_idx']]
            para.clear()
            self._apply_runs(para, segment['runs_list'])
            
            if progress_callback:
                progress_callback()
    
    # @progress_tracker(item_name='chart segments', use_tqdm=True)
    def _inject_xml_content(self, segments: list[ChartSegment | SmartArtSegment], temp_dir: str, inject_func, progress_callback=None):
        """Hàm chung để inject nội dung vào các file XML"""
        count = 0
        for seg in segments:
            try:
                file_path = os.path.join(temp_dir, seg['file_path'])
                if not os.path.exists(file_path):
                    if progress_callback:
                        progress_callback()
                    continue
                
                tree = etree.parse(file_path)
                root = tree.getroot()
                
                if inject_func(root, seg):
                    count += 1
                    tree.write(file_path, encoding='UTF-8', xml_declaration=True, pretty_print=False)
            except Exception as e:
                self.logger.warning(f"   ⚠️  Error injecting segment: {e}")
            finally:
                if progress_callback:
                    progress_callback()
        
        return count
    
    def _inject_chart_element(self, root, seg: ChartSegment):
        """Inject text vào chart element"""
        if seg['element_type'] == 'title':
            titles = root.findall('.//c:title', self.ns)
            if seg['element_idx'] < len(titles):
                text_elems = titles[seg['element_idx']].findall('.//a:t', self.ns)
                if text_elems:
                    text_elems[0].text = seg['translated_text']
                    return True
        
        elif seg['element_type'] == 'value':
            v_elements = root.findall('.//c:v', self.ns)
            count = 0
            for v_elem in v_elements:
                if v_elem.text and v_elem.text.strip():
                    text = v_elem.text.strip()
                    if not is_numeric(text):
                        if count == seg['element_idx']:
                            v_elem.text = seg['translated_text']
                            return True
                        count += 1
        return False
    
    def _inject_smartart_element(self, root, seg: SmartArtSegment):
        """Inject text vào SmartArt element"""
        text_elems = root.findall('.//a:t', self.ns)
        if seg['element_idx'] < len(text_elems):
            text_elems[seg['element_idx']].text = seg['translated_text']
            return True
        return False
    
    def _inject_chart_and_smartart(self, chart_segments: list[ChartSegment], smartart_segments: list[SmartArtSegment]):
        """Chèn văn bản vào biểu đồ và SmartArt bằng cách sửa trực tiếp file XML"""
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Giải nén tất cả các file
            with zipfile.ZipFile(self.output_file, 'r') as z:
                z.extractall(temp_dir)
            
            # Inject chart segments với progress bar
            self._inject_xml_content(chart_segments, temp_dir, self._inject_chart_element)
            
            # Inject SmartArt segments với progress bar
            self._inject_xml_content(smartart_segments, temp_dir, self._inject_smartart_element)
            
            # Đóng gói lại thành file ZIP
            with zipfile.ZipFile(self.output_file, 'w', zipfile.ZIP_DEFLATED) as z:
                for root_dir, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root_dir, file)
                        arc_path = os.path.relpath(file_path, temp_dir)
                        z.write(file_path, arc_path)
        
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    @timer
    @log_errors
    def inject(self):
        """Đọc checkpoint và chèn nội dung đã dịch vào tài liệu"""
        self.logger.info("="*70)
        self.logger.info("INJECTING TRANSLATIONS BACK TO DOCX")
        self.logger.info("="*70 + "\n")
        
        with open(self.checkpoint_file, "r", encoding="utf-8") as f:
            checkpoint_data = json.load(f)
        
        # Inject text và table segments với progress bar
        self._inject_text_segments(checkpoint_data["text_segments"])
        self._inject_table_cell_segments(checkpoint_data["table_cell_segments"])
        
        # Lưu trước khi inject chart/smartart
        self.doc.save(self.output_file)
        
        # Inject chart và SmartArt segments với progress bar
        self._inject_chart_and_smartart(
            checkpoint_data["chart_segments"],
            checkpoint_data["smartart_segments"]
        )
        
        self.logger.info(f"☑️ Output saved to {self.output_file}")
        os.remove(self.checkpoint_file)
        self.logger.info(f"☑️ Checkpoint file deleted: {self.checkpoint_file}")