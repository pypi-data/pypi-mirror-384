"""
Word文档内容添加工具函数

提供五个核心的文档内容添加功能
"""

import os
from typing import Optional, List
from .utils import (
    ensure_docx_extension,
    check_file_writeable
)


async def create_document(filename: str, template: Optional[str] = None, overwrite: bool = False) -> str:
    """Create a new Word (.docx) document.

    Args:
        filename: Path to the target Word document ('.docx' will be appended if missing)
        template: Optional path to a template .docx file to base the new document on
        overwrite: Allow overwriting the file if it already exists (default False)
    """
    from docx import Document

    # Ensure .docx extension
    filename = ensure_docx_extension(filename)

    # Ensure target directory exists (if any directory part present)
    try:
        dir_name = os.path.dirname(filename)
        if dir_name:
            os.makedirs(dir_name, exist_ok=True)
    except Exception:
        # Non-fatal: directory creation failure will be surfaced by save()
        pass

    # Handle overwrite / existence
    if os.path.exists(filename):
        if not overwrite:
            raise FileExistsError(f"Target document already exists: {os.path.abspath(filename)}. Set overwrite=True to replace it.")
        # Check writability for overwrite
        is_writeable, error_message = check_file_writeable(filename)
        if not is_writeable:
            raise PermissionError(f"Cannot overwrite document: {error_message}. Close the file or choose a different path.")

    # Create document (from template or blank)
    try:
        if template:
            abs_template = os.path.abspath(template)
            if not os.path.exists(abs_template):
                raise FileNotFoundError(f"Template file not found: {abs_template}")
            try:
                doc = Document(abs_template)
            except Exception as open_error:
                try:
                    from docx.opc.exceptions import PackageNotFoundError
                except Exception:
                    PackageNotFoundError = tuple()  # type: ignore
                if isinstance(open_error, PackageNotFoundError):
                    raise RuntimeError(
                        f"Template is not a valid .docx package or is corrupted: {abs_template}."
                    )
                raise
        else:
            doc = Document()

        # Save to target path
        abs_target = os.path.abspath(filename)
        doc.save(abs_target)
        return f"Document created at {abs_target}" + (f" using template {os.path.abspath(template)}" if template else "")
    except Exception as e:
        raise RuntimeError(f"Failed to create document: {str(e)}")

async def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    
    filename = ensure_docx_extension(filename)
    
    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: level must be an integer between 1 and 9")
    
    # Validate level range
    if level < 1 or level > 9:
        raise ValueError(f"Invalid heading level: {level}. Level must be between 1 and 9.")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document.")
    
    try:
        # 捕获无效docx包的错误，提供更清晰提示
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        # 统一使用段落样式方式，并在缺少样式时创建；先尝试升级末段，避免重复插入
        style_name = 'Heading ' + str(level)
        try:
            _ = doc.styles[style_name]
        except KeyError:
            # 创建基本标题样式（英文名称），以适配不同本地化模板
            new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            try:
                new_style.base_style = doc.styles['Normal']
            except Exception:
                pass
            new_style.font.bold = True
            if level == 1:
                new_style.font.size = Pt(16)
            elif level == 2:
                new_style.font.size = Pt(14)
            else:
                new_style.font.size = Pt(12)

        # 如果文档末段文本与将插入的标题相同，则直接将该段落升级为标题，避免“正文+标题”重复
        if doc.paragraphs and doc.paragraphs[-1].text.strip() == str(text).strip():
            paragraph = doc.paragraphs[-1]
            try:
                paragraph.style = style_name
                doc.save(filename)
                return f"Upgraded last paragraph to heading '{text}' (level {level}) in {filename}"
            except Exception:
                # 若设置样式失败，则进行直接格式化作为兜底
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(16)
                    elif level == 2:
                        run.font.size = Pt(14)
                    else:
                        run.font.size = Pt(12)
                doc.save(filename)
                return f"Upgraded last paragraph with direct formatting for heading '{text}' in {filename}"

        # 正常路径：插入一个标题段落（使用样式）
        paragraph = doc.add_paragraph(text)
        try:
            paragraph.style = style_name
        except Exception:
            # 样式设置失败则直接格式化
            if paragraph.runs:
                run = paragraph.runs[0]
                run.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        raise RuntimeError(f"Failed to add heading: {str(e)}")


async def add_paragraph(filename: str, text: str, style: Optional[str] = None) -> str:
    """Add a paragraph to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        paragraph = doc.add_paragraph(text)
        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"
        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        raise RuntimeError(f"Failed to add paragraph: {str(e)}")


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        table = doc.add_table(rows=rows, cols=cols)
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        raise RuntimeError(f"Failed to add table: {str(e)}")


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.

    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    from docx import Document
    from docx.shared import Inches

    filename = ensure_docx_extension(filename)

    # Validate document existence
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)

    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        raise FileNotFoundError(f"Image file not found: {abs_image_path}")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(abs_filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {abs_filename}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        diagnostic = f"Document: {abs_filename}, Image: {abs_image_path}"
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            raise RuntimeError(f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}")
    except Exception as outer_error:
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        raise RuntimeError(f"Document processing error: {error_type} - {error_msg or 'No error details available'}")


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.

    Args:
        filename: Path to the Word document
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        raise RuntimeError(f"Failed to add page break: {str(e)}")
