"""
Document Processing Module
Handles Word to Markdown and Markdown to Word conversions
"""

import os
import asyncio
from pathlib import Path
from typing import Dict, Any, Optional
import logging

import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document format conversions"""
    
    def __init__(self):
        self.temp_dir = Path(os.getenv('TEMP_DIR', './temp_documents'))
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    async def word_to_markdown(self, word_path: str) -> Dict[str, Any]:
        """
        Convert Word document to Markdown
        
        Args:
            word_path: Path to the Word document
            
        Returns:
            Dict with converted content and metadata
        """
        try:
            word_file = Path(word_path)
            if not word_file.exists():
                raise FileNotFoundError(f"Word document not found: {word_path}")
            
            # Validate file size
            if word_file.stat().st_size == 0:
                raise ValueError(f"File is empty: {word_path}")
            
            # Validate it's a valid ZIP file (docx is a ZIP archive)
            try:
                import zipfile
                with zipfile.ZipFile(str(word_file.absolute()), 'r') as zf:
                    # Check for required docx structure
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    missing_files = [f for f in required_files if f not in zf.namelist()]
                    if missing_files:
                        raise ValueError(f"Invalid docx structure, missing: {missing_files}")
                    logger.info(f"File validation passed for: {word_path}")
            except zipfile.BadZipFile:
                raise ValueError(f"File is not a valid ZIP/docx archive: {word_path}")
            except ValueError:
                raise  # Re-raise validation errors
            except Exception as e:
                logger.warning(f"ZIP validation warning for {word_path}: {e}")
            
            # Convert to absolute path and handle special characters
            abs_path = word_file.resolve()
            
            # Read Word document using different approaches for problematic filenames
            doc = None
            try:
                # First try with the absolute path
                doc = docx.Document(str(abs_path))
            except Exception as e1:
                logger.warning(f"Failed to load with absolute path: {e1}")
                try:
                    # Try with original path
                    doc = docx.Document(word_path)
                except Exception as e2:
                    logger.warning(f"Failed to load with original path: {e2}")
                    try:
                        # Try reading as binary and passing to docx
                        with open(abs_path, 'rb') as f:
                            doc = docx.Document(f)
                    except Exception as e3:
                        # If all methods fail, raise the original error with context
                        raise Exception(f"Unable to load Word document '{word_path}'. Tried multiple approaches. Last error: {e3}")
            
            if doc is None:
                raise Exception(f"Failed to load Word document: {word_path}")
            
            # Extract content and convert to markdown
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Detect and format headers based on style
                style_name = paragraph.style.name.lower()
                if 'heading 1' in style_name or paragraph.style.name == 'Heading 1':
                    markdown_content.append(f"# {text}\n")
                elif 'heading 2' in style_name or paragraph.style.name == 'Heading 2':
                    markdown_content.append(f"## {text}\n")
                elif 'heading 3' in style_name or paragraph.style.name == 'Heading 3':
                    markdown_content.append(f"### {text}\n")
                elif 'heading 4' in style_name or paragraph.style.name == 'Heading 4':
                    markdown_content.append(f"#### {text}\n")
                else:
                    # Regular paragraph
                    # Check for bold/italic formatting
                    formatted_text = self._process_paragraph_formatting(paragraph)
                    markdown_content.append(f"{formatted_text}\n\n")
            
            # Process tables if any
            for table in doc.tables:
                table_md = self._convert_table_to_markdown(table)
                markdown_content.append(f"{table_md}\n\n")
            
            # Join all content
            full_content = "".join(markdown_content)
            
            # Save to temp file
            output_path = self.temp_dir / f"{word_file.stem}.md"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            logger.info(f"Successfully converted Word to Markdown: {output_path}")
            
            return {
                "content": full_content,
                "output_path": str(output_path),
                "word_count": len(full_content.split()),
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"Error converting Word to Markdown: {e}")
            raise
    
    def _process_paragraph_formatting(self, paragraph) -> str:
        """Process paragraph with inline formatting"""
        text = ""
        for run in paragraph.runs:
            run_text = run.text
            
            # Apply formatting
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
                
            text += run_text
        
        return text
    
    def _convert_table_to_markdown(self, table) -> str:
        """Convert Word table to Markdown table"""
        rows = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            row_text = "| " + " | ".join(cells) + " |"
            rows.append(row_text)
            
            # Add separator after header row
            if i == 0:
                separator = "|" + "|".join([" --- " for _ in cells]) + "|"
                rows.append(separator)
        
        return "\n".join(rows)
    
    async def markdown_to_word(self, markdown_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Convert Markdown document to Word
        
        Args:
            markdown_path: Path to the Markdown file
            output_path: Optional output path for Word document
            
        Returns:
            Dict with conversion results
        """
        try:
            markdown_file = Path(markdown_path)
            if not markdown_file.exists():
                raise FileNotFoundError(f"Markdown file not found: {markdown_path}")
            
            # Read Markdown content
            with open(markdown_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create new Word document
            doc = docx.Document()
            
            # Split content by lines for processing
            lines = content.split('\n')
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                # Process headers
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    heading = doc.add_heading(line[5:], level=4)
                
                # Process tables
                elif line.startswith('|'):
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        if '---' not in lines[i]:  # Skip separator row
                            table_lines.append(lines[i].strip())
                        i += 1
                    
                    if table_lines:
                        self._create_word_table(doc, table_lines)
                    i -= 1  # Adjust for the extra increment
                
                # Process regular paragraphs
                else:
                    # Collect paragraph lines
                    paragraph_lines = []
                    while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('|'):
                        paragraph_lines.append(lines[i].strip())
                        i += 1
                    
                    if paragraph_lines:
                        paragraph_text = ' '.join(paragraph_lines)
                        # Process inline formatting and add paragraph
                        self._add_formatted_paragraph(doc, paragraph_text)
                    i -= 1  # Adjust for the extra increment
                
                i += 1
            
            # Set output path
            if not output_path:
                output_path = self.temp_dir / f"{markdown_file.stem}.docx"
            else:
                output_path = Path(output_path)
            
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Save document
            doc.save(str(output_path))
            
            # Get document statistics
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)
            
            # Estimate word count
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())
            
            logger.info(f"Successfully converted Markdown to Word: {output_path}")
            
            return {
                "output_path": str(output_path),
                "word_count": word_count,
                "paragraph_count": total_paragraphs,
                "table_count": total_tables
            }
            
        except Exception as e:
            logger.error(f"Error converting Markdown to Word: {e}")
            raise
    
    def _create_word_table(self, doc, table_lines):
        """Create a Word table from markdown table lines"""
        if not table_lines:
            return
        
        # Parse table data
        rows_data = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
            if cells:
                rows_data.append(cells)
        
        if not rows_data:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        table.style = 'Table Grid'
        
        # Fill table data
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_text
    
    def _add_formatted_paragraph(self, doc, text):
        """Add a paragraph with inline formatting to Word document"""
        paragraph = doc.add_paragraph()
        
        # Simple regex patterns for formatting
        parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if part.startswith('***') and part.endswith('***'):
                # Bold and italic
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # Regular text
                paragraph.add_run(part)
    
    async def get_document_info(self, document_path: str) -> Dict[str, Any]:
        """
        Get information about a document
        
        Args:
            document_path: Path to the document
            
        Returns:
            Dictionary with document information
        """
        try:
            doc_path = Path(document_path)
            if not doc_path.exists():
                raise FileNotFoundError(f"Document not found: {document_path}")
            
            info = {
                "name": doc_path.name,
                "path": str(doc_path),
                "size": doc_path.stat().st_size,
                "extension": doc_path.suffix.lower()
            }
            
            if doc_path.suffix.lower() == '.docx':
                doc = docx.Document(str(doc_path))
                info.update({
                    "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "table_count": len(doc.tables),
                    "estimated_word_count": sum(len(p.text.split()) for p in doc.paragraphs)
                })
            elif doc_path.suffix.lower() == '.md':
                with open(doc_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                info.update({
                    "line_count": len(content.split('\n')),
                    "character_count": len(content),
                    "word_count": len(content.split())
                })
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting document info: {e}")
            raise
