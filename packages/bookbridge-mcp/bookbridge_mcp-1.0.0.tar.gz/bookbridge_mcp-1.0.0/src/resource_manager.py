"""
Resource Management Module
Manages document resources and file operations
"""

import os
import asyncio
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging
import json
from datetime import datetime

logger = logging.getLogger(__name__)


class ResourceManager:
    """Manages document resources and file operations"""
    
    def __init__(self):
        self.input_dir = Path(os.getenv('INPUT_DIR', './input_documents'))
        self.output_dir = Path(os.getenv('OUTPUT_DIR', './output_documents'))
        self.temp_dir = Path(os.getenv('TEMP_DIR', './temp_documents'))
        
        # Create directories if they don't exist
        for dir_path in [self.input_dir, self.output_dir, self.temp_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)
        
        # Document registry file
        self.registry_file = self.temp_dir / "document_registry.json"
        self.registry = self._load_registry()
    
    def _load_registry(self) -> Dict[str, Any]:
        """Load document registry from file"""
        try:
            if self.registry_file.exists():
                with open(self.registry_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                return {
                    "documents": {},
                    "last_updated": datetime.now().isoformat()
                }
        except Exception as e:
            logger.error(f"Error loading registry: {e}")
            return {
                "documents": {},
                "last_updated": datetime.now().isoformat()
            }
    
    def _save_registry(self):
        """Save document registry to file"""
        try:
            self.registry["last_updated"] = datetime.now().isoformat()
            with open(self.registry_file, 'w', encoding='utf-8') as f:
                json.dump(self.registry, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving registry: {e}")
    
    async def get_source_document(self, document_id: str) -> str:
        """
        Get source document content by ID
        
        Args:
            document_id: Document identifier
            
        Returns:
            Document content as string
        """
        try:
            # First check registry
            if document_id in self.registry["documents"]:
                doc_info = self.registry["documents"][document_id]
                file_path = Path(doc_info["source_path"])
                
                if file_path.exists():
                    if file_path.suffix.lower() == '.docx':
                        # Read Word document content
                        import docx
                        doc = docx.Document(str(file_path))
                        content_parts = []
                        
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                content_parts.append(paragraph.text)
                        
                        return "\n\n".join(content_parts)
                    
                    elif file_path.suffix.lower() == '.md':
                        # Read Markdown content
                        with open(file_path, 'r', encoding='utf-8') as f:
                            return f.read()
            
            # If not in registry, try to find by filename
            source_files = list(self.input_dir.glob(f"*{document_id}*"))
            if source_files:
                return await self.get_source_document_by_path(str(source_files[0]))
            
            raise FileNotFoundError(f"Source document not found: {document_id}")
            
        except Exception as e:
            logger.error(f"Error retrieving source document {document_id}: {e}")
            raise
    
    async def get_source_document_by_path(self, file_path: str) -> str:
        """Get source document content by file path"""
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            if path.suffix.lower() == '.docx':
                import docx
                doc = docx.Document(str(path))
                content_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                
                return "\n\n".join(content_parts)
            
            elif path.suffix.lower() in ['.md', '.txt']:
                with open(path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            else:
                raise ValueError(f"Unsupported file format: {path.suffix}")
                
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            raise
    
    async def get_intermediate_document(self, document_id: str) -> str:
        """
        Get intermediate Markdown document content
        
        Args:
            document_id: Document identifier
            
        Returns:
            Markdown content as string
        """
        try:
            # Check registry
            if document_id in self.registry["documents"]:
                doc_info = self.registry["documents"][document_id]
                if "intermediate_path" in doc_info:
                    file_path = Path(doc_info["intermediate_path"])
                    if file_path.exists():
                        with open(file_path, 'r', encoding='utf-8') as f:
                            return f.read()
            
            # Try to find in temp directory
            markdown_files = list(self.temp_dir.glob(f"*{document_id}*.md"))
            if markdown_files:
                with open(markdown_files[0], 'r', encoding='utf-8') as f:
                    return f.read()
            
            raise FileNotFoundError(f"Intermediate document not found: {document_id}")
            
        except Exception as e:
            logger.error(f"Error retrieving intermediate document {document_id}: {e}")
            raise
    
    async def get_target_document(self, document_id: str) -> str:
        """
        Get target document content
        
        Args:
            document_id: Document identifier
            
        Returns:
            Document content as string
        """
        try:
            # Check registry
            if document_id in self.registry["documents"]:
                doc_info = self.registry["documents"][document_id]
                if "target_path" in doc_info:
                    file_path = Path(doc_info["target_path"])
                    if file_path.exists():
                        if file_path.suffix.lower() == '.docx':
                            import docx
                            doc = docx.Document(str(file_path))
                            content_parts = []
                            
                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():
                                    content_parts.append(paragraph.text)
                            
                            return "\n\n".join(content_parts)
                        
                        elif file_path.suffix.lower() == '.md':
                            with open(file_path, 'r', encoding='utf-8') as f:
                                return f.read()
            
            # Try to find in output directory
            target_files = list(self.output_dir.glob(f"*{document_id}*"))
            if target_files:
                return await self.get_source_document_by_path(str(target_files[0]))
            
            raise FileNotFoundError(f"Target document not found: {document_id}")
            
        except Exception as e:
            logger.error(f"Error retrieving target document {document_id}: {e}")
            raise
    
    async def register_document(
        self,
        document_id: str,
        source_path: str,
        intermediate_path: Optional[str] = None,
        target_path: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ):
        """
        Register a document in the resource manager
        
        Args:
            document_id: Unique document identifier
            source_path: Path to source document
            intermediate_path: Optional path to intermediate document
            target_path: Optional path to target document
            metadata: Optional metadata dictionary
        """
        try:
            doc_info = {
                "document_id": document_id,
                "source_path": str(source_path),
                "created_at": datetime.now().isoformat(),
                "status": "registered"
            }
            
            if intermediate_path:
                doc_info["intermediate_path"] = str(intermediate_path)
            
            if target_path:
                doc_info["target_path"] = str(target_path)
            
            if metadata:
                doc_info["metadata"] = metadata
            
            self.registry["documents"][document_id] = doc_info
            self._save_registry()
            
            logger.info(f"Document registered: {document_id}")
            
        except Exception as e:
            logger.error(f"Error registering document: {e}")
            raise
    
    async def update_document_status(
        self,
        document_id: str,
        status: str,
        intermediate_path: Optional[str] = None,
        target_path: Optional[str] = None
    ):
        """Update document processing status"""
        try:
            if document_id in self.registry["documents"]:
                doc_info = self.registry["documents"][document_id]
                doc_info["status"] = status
                doc_info["updated_at"] = datetime.now().isoformat()
                
                if intermediate_path:
                    doc_info["intermediate_path"] = str(intermediate_path)
                
                if target_path:
                    doc_info["target_path"] = str(target_path)
                
                self._save_registry()
                logger.info(f"Document status updated: {document_id} -> {status}")
            
        except Exception as e:
            logger.error(f"Error updating document status: {e}")
            raise
    
    async def list_all_documents(self) -> Dict[str, List[Dict[str, Any]]]:
        """
        List all available documents
        
        Returns:
            Dictionary with categorized document listings
        """
        try:
            # Scan directories for files
            source_files = self._scan_directory(self.input_dir, ['*.docx', '*.doc', '*.md', '*.txt'])
            intermediate_files = self._scan_directory(self.temp_dir, ['*.md'])
            target_files = self._scan_directory(self.output_dir, ['*.docx', '*.doc', '*.md'])
            
            # Get registered documents
            registered_docs = []
            for doc_id, doc_info in self.registry["documents"].items():
                registered_docs.append({
                    "id": doc_id,
                    "status": doc_info.get("status", "unknown"),
                    "source_path": doc_info.get("source_path", ""),
                    "intermediate_path": doc_info.get("intermediate_path", ""),
                    "target_path": doc_info.get("target_path", ""),
                    "created_at": doc_info.get("created_at", ""),
                    "updated_at": doc_info.get("updated_at", "")
                })
            
            return {
                "registered_documents": registered_docs,
                "source_files": source_files,
                "intermediate_files": intermediate_files,
                "target_files": target_files,
                "total_registered": len(registered_docs),
                "total_source": len(source_files),
                "total_intermediate": len(intermediate_files),
                "total_target": len(target_files)
            }
            
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            raise
    
    def _scan_directory(self, directory: Path, patterns: List[str]) -> List[Dict[str, Any]]:
        """Scan directory for files matching patterns"""
        files = []
        try:
            for pattern in patterns:
                for file_path in directory.glob(pattern):
                    if file_path.is_file():
                        stat_info = file_path.stat()
                        files.append({
                            "name": file_path.name,
                            "path": str(file_path),
                            "size": stat_info.st_size,
                            "modified": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
                            "extension": file_path.suffix.lower()
                        })
        except Exception as e:
            logger.error(f"Error scanning directory {directory}: {e}")
        
        return sorted(files, key=lambda x: x["modified"], reverse=True)
    
    async def cleanup_temp_files(self, older_than_hours: int = 24):
        """Clean up temporary files older than specified hours"""
        try:
            from datetime import timedelta
            
            cutoff_time = datetime.now() - timedelta(hours=older_than_hours)
            cleaned_files = []
            
            for file_path in self.temp_dir.glob("*"):
                if file_path.is_file():
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_time < cutoff_time:
                        try:
                            file_path.unlink()
                            cleaned_files.append(str(file_path))
                        except Exception as e:
                            logger.error(f"Error deleting {file_path}: {e}")
            
            if cleaned_files:
                logger.info(f"Cleaned up {len(cleaned_files)} temporary files")
            
            return {
                "cleaned_files": len(cleaned_files),
                "files": cleaned_files
            }
            
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")
            raise
    
    async def get_document_stats(self) -> Dict[str, Any]:
        """Get overall document statistics"""
        try:
            all_docs = await self.list_all_documents()
            
            # Calculate status distribution
            status_counts = {}
            for doc in all_docs["registered_documents"]:
                status = doc["status"]
                status_counts[status] = status_counts.get(status, 0) + 1
            
            # Calculate total file sizes
            total_size = 0
            for category in ["source_files", "intermediate_files", "target_files"]:
                for file_info in all_docs[category]:
                    total_size += file_info["size"]
            
            return {
                "total_registered_documents": all_docs["total_registered"],
                "total_source_files": all_docs["total_source"],
                "total_intermediate_files": all_docs["total_intermediate"],
                "total_target_files": all_docs["total_target"],
                "status_distribution": status_counts,
                "total_storage_bytes": total_size,
                "total_storage_mb": round(total_size / (1024 * 1024), 2),
                "registry_location": str(self.registry_file),
                "directories": {
                    "input": str(self.input_dir),
                    "output": str(self.output_dir),
                    "temp": str(self.temp_dir)
                }
            }
            
        except Exception as e:
            logger.error(f"Error getting document stats: {e}")
            raise
