from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config
import re

class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = self.format_markdown(markdown_content)
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        self.html_content = self.move_p_content_to_li(self.html_content)
        FileUtils.create_paths(Config.BASE_PATH)

    def format_markdown(markdown_text):
        """
        逐行处理Markdown格式，确保格式正确
        
        Args:
            markdown_text (str): 原始Markdown文本
            
        Returns:
            str: 格式化后的Markdown文本
        """
        lines = markdown_text.split('\n')
        processed_lines = []
        
        for i, line in enumerate(lines):
            # 处理标题行：去掉前面的空格，确保#号后面有空格
            if line.strip().startswith('#'):
                line = re.sub(r'^\s*(#+)\s*', r'\1 ', line)
            
            # 处理表格和标题之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and line.strip().startswith('#'):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理表格和列表之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and re.match(r'^\s*(\d+\.|\-)', line.strip()):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理冒号后的换行
            line = line.replace("：\n", "：\n\n")
            
            # 处理列表项格式：确保列表项有正确的缩进
            if re.match(r'^\s*\d+\.', line.strip()):
                # 数字列表项
                line = re.sub(r'^(\s*)(\d+\.\s+)', r'\1\2', line)
            elif re.match(r'^\s*\-', line.strip()):
                # 项目符号列表项
                line = re.sub(r'^(\s*)(\-\s+)', r'\1\2', line)
            
            # 处理列表项之间的空行
            if i > 0 and re.match(r'^\s*\d+\.', lines[i-1].strip()) and re.match(r'^\s*\d+\.', line.strip()):
                # 连续的数字列表项之间不添加空行
                pass
            elif i > 0 and re.match(r'^\s*\-', lines[i-1].strip()) and re.match(r'^\s*\-', line.strip()):
                # 连续的项目符号列表项之间不添加空行
                pass
            
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)

    def to_html(self):
        return self.html_content

    def move_p_content_to_li(self,html):
        soup = BeautifulSoup(html, 'html.parser')

        for ul in soup.find_all('ul'):
            for li in ul.find_all('li'):
                for p in li.find_all('p'):
                    p_content = "".join(p.decode_contents())
                    p.insert_before(BeautifulSoup(p_content, 'html.parser'))
                    p.decompose()
                for child in li.contents:
                    if isinstance(child, NavigableString) and child.strip() == '':
                        child.extract()

        return str(soup)
    def to_word(self,file_name):
        """
        将markdown文本转换成word，
        :param file_name
        :return: 返回文件名
        """
        if file_name is None or file_name =='':
            file_name = CommonUtils.generate_uuid() + ".docx"
        elif not file_name.endswith(".docx"):
            file_name += ".docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img','ul','li']):
            if tag.name == 'h1':
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h5':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h6':
                heading = doc.add_heading(level=4)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'p':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(tag.get_text())
                run.font.size = Pt(12)
            elif tag.name == 'img':
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        fixed_width = Inches(6)

                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'ul':
                for li in tag.find_all('li'):
                    paragraph = doc.add_paragraph(style='ListBullet')
                    for child in li.children:
                        if child.name == 'strong':
                            run = paragraph.add_run(child.get_text())
                            run.bold = True
                        else:
                            paragraph.add_run(child)
            elif tag.name == 'table':
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for cell in table.rows[0].cells:
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1

                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,file_name,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param file_name:
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        if file_name is None or file_name=='':
            file_name = CommonUtils.generate_uuid()+ ".pdf";
        elif not file_name.endswith(".pdf"):
            file_name += ".pdf"
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name
