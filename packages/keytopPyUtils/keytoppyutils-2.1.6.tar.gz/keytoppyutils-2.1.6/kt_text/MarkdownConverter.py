from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config
import re

class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = self.format_markdown(markdown_content)
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        self.html_content = self.move_p_content_to_li(self.html_content)
        FileUtils.create_paths(Config.BASE_PATH)

    def format_markdown(self, markdown_text):
        """
        逐行处理Markdown格式，确保格式正确，特别是嵌套列表结构
        
        Args:
            markdown_text (str): 原始Markdown文本
            
        Returns:
            str: 格式化后的Markdown文本
        """
        lines = markdown_text.split('\n')
        processed_lines = []
        
        # 跟踪当前缩进级别
        current_indent = 0
        
        for i, line in enumerate(lines):
            line = line.rstrip()  # 移除行尾空白
            
            # 跳过空行
            if not line.strip():
                processed_lines.append('')
                continue
            
            # 处理标题行：去掉前面的空格，确保#号后面有空格
            if line.strip().startswith('#'):
                line = re.sub(r'^\s*(#+)\s*', r'\1 ', line)
                current_indent = 0  # 标题重置缩进
            
            # 处理列表项
            elif re.match(r'^\s*(\d+\.|\-)', line.strip()):
                # 计算当前缩进级别
                indent_level = len(re.match(r'^(\s*)', line).group(1)) // 2  # 每2个空格一个级别
                
                # 确保缩进级别合理
                if indent_level > current_indent + 1:
                    indent_level = current_indent + 1
                
                # 更新当前缩进级别
                current_indent = indent_level
                
                # 格式化列表项：确保正确的缩进
                indent_spaces = '  ' * indent_level
                
                # 处理数字列表项
                if re.match(r'^\s*\d+\.', line.strip()):
                    line = re.sub(r'^(\s*)(\d+\.\s+)', f'{indent_spaces}\2', line)
                # 处理项目符号列表项
                elif re.match(r'^\s*\-', line.strip()):
                    line = re.sub(r'^(\s*)(\-\s+)', f'{indent_spaces}\2', line)
            
            # 处理表格和标题之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and line.strip().startswith('#'):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理表格和列表之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and re.match(r'^\s*(\d+\.|\-)', line.strip()):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理冒号后的换行
            line = line.replace("：\n", "：\n\n")
            
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)

    def move_p_content_to_li(self, html):
        """修复HTML中的嵌套列表结构"""
        soup = BeautifulSoup(html, 'html.parser')
        
        # 修复嵌套列表结构
        self._fix_nested_list_structure(soup)
        
        # 处理所有列表类型：ul和ol
        for list_tag in soup.find_all(['ul', 'ol']):
            for li in list_tag.find_all('li', recursive=False):
                self._process_li_html_content(li)

        return str(soup)

    def _fix_nested_list_structure(self, soup):
        """修复嵌套列表结构，确保ul/ol在正确的li内部"""
        # 首先修复嵌套列表的层级关系
        self._fix_nested_list_hierarchy(soup)
        
        # 然后修复列表项的结构
        self._fix_li_structure(soup)

    def _fix_li_structure(self, soup):
        """修复列表项结构，确保同一级别的列表项不被错误嵌套"""
        # 查找所有列表
        for list_tag in soup.find_all(['ul', 'ol']):
            # 获取所有直接子列表项
            li_items = list_tag.find_all('li', recursive=False)
            
            for li in li_items:
                # 检查列表项内容
                li_text = li.get_text(strip=True)
                
                # 如果列表项包含嵌套列表，确保结构正确
                nested_lists = li.find_all(['ul', 'ol'], recursive=False)
                
                for nested_list in nested_lists:
                    # 检查嵌套列表是否应该属于当前列表项
                    if not self._should_nest_under_current_li(li, nested_list):
                        # 如果不应该嵌套，将嵌套列表移到当前列表后面
                        li.insert_after(nested_list.extract())

    def _should_nest_under_current_li(self, li, nested_list):
        """判断嵌套列表是否应该属于当前列表项"""
        # 获取当前列表项的内容
        li_text = li.get_text(strip=True)
        
        # 如果列表项以"计费规则"、"收费标准"等关键词结尾，应该包含嵌套列表
        nesting_keywords = ['计费规则', '收费标准', '收费规则', '规则']
        
        for keyword in nesting_keywords:
            if keyword in li_text:
                return True
        
        # 默认情况下，如果嵌套列表紧跟在列表项文本后面，应该属于该列表项
        # 检查嵌套列表前面的内容
        prev_sibling = nested_list.previous_sibling
        if prev_sibling and isinstance(prev_sibling, NavigableString) and prev_sibling.strip():
            return True
        
        return False

    def get_html_structure(self):
        """获取HTML结构分析，用于调试"""
        soup = BeautifulSoup(self.html_content, 'html.parser')
        
        result = []
        result.append("=== HTML结构分析 ===")
        
        # 检查所有列表
        lists = soup.find_all(['ul', 'ol'])
        for i, list_tag in enumerate(lists):
            result.append(f"\n列表 {i+1} ({list_tag.name}):")
            
            # 检查列表的父元素
            parent = list_tag.parent
            result.append(f"  父元素: {parent.name if parent else 'None'}")
            
            # 检查列表项
            li_items = list_tag.find_all('li', recursive=False)
            result.append(f"  直接子列表项数量: {len(li_items)}")
            
            for j, li in enumerate(li_items):
                result.append(f"  列表项 {j+1}:")
                # 获取列表项内容
                li_text = li.get_text(strip=True)
                result.append(f"    内容: {li_text[:50]}..." if len(li_text) > 50 else f"    内容: {li_text}")
                
                # 检查是否有嵌套列表
                nested_lists = li.find_all(['ul', 'ol'], recursive=False)
                result.append(f"    嵌套列表数量: {len(nested_lists)}")
                
                for k, nested_list in enumerate(nested_lists):
                    nested_parent = nested_list.parent
                    result.append(f"      嵌套列表 {k+1} 父元素: {nested_parent.name if nested_parent else 'None'}")
        
        return '\n'.join(result)

    def debug_html_content(self):
        """调试方法：打印HTML内容和结构"""
        print("=== 原始HTML内容 ===")
        print(self.html_content)
        print("\n" + "="*50 + "\n")
        
        print("=== HTML结构分析 ===")
        print(self.get_html_structure())
        # 查找所有li标签
        for li in soup.find_all('li'):
            # 检查li的内容是否包含ul/ol标签
            for child in li.contents:
                if hasattr(child, 'name') and child.name in ['ul', 'ol']:
                    # 如果ul/ol是li的直接子元素，结构正确
                    continue
                
                # 检查文本内容后是否应该跟ul/ol
                if isinstance(child, NavigableString) and child.strip():
                    # 查找文本内容后的ul/ol标签
                    next_sibling = child.next_sibling
                    if next_sibling and hasattr(next_sibling, 'name') and next_sibling.name in ['ul', 'ol']:
                        # 将ul/ol移动到当前li内部
                        child.insert_after(next_sibling.extract())
        
        # 修复空的列表项
        for li in soup.find_all('li'):
            if not li.contents or (len(li.contents) == 1 and isinstance(li.contents[0], NavigableString) and not li.contents[0].strip()):
                # 删除空的列表项
                li.decompose()
        
        # 修复嵌套列表的层级关系
        self._fix_nested_list_hierarchy(soup)

    def _fix_nested_list_hierarchy(self, soup):
        """修复嵌套列表的层级关系"""
        # 查找所有ul和ol标签
        for list_tag in soup.find_all(['ul', 'ol']):
            # 检查列表是否在正确的li内部
            parent = list_tag.parent
            if parent and parent.name != 'li':
                # 如果列表不在li内部，需要找到最近的li祖先
                li_ancestor = list_tag.find_parent('li')
                if li_ancestor:
                    # 将列表移动到li祖先内部
                    li_ancestor.append(list_tag.extract())
                else:
                    # 如果没有li祖先，创建一个新的li包装
                    new_li = soup.new_tag('li')
                    list_tag.wrap(new_li)
    def to_html(self):
        return self.html_content

    def _process_li_html_content(self, li):
        """递归处理li标签内容（HTML预处理）"""
        # 处理当前li中的p标签
        for p in li.find_all('p'):
            p_content = "".join(p.decode_contents())
            p.insert_before(BeautifulSoup(p_content, 'html.parser'))
            p.decompose()
        
        # 移除空文本节点
        for child in li.contents:
            if isinstance(child, NavigableString) and child.strip() == '':
                child.extract()
        
        # 递归处理嵌套的列表
        for nested_list in li.find_all(['ul', 'ol'], recursive=False):
            for nested_li in nested_list.find_all('li', recursive=False):
                self._process_li_html_content(nested_li)

    def _process_list(self, doc, list_tag, list_style, level=0):
        """处理列表标签，支持嵌套列表"""
        for li in list_tag.find_all('li', recursive=False):
            # 创建列表项段落
            paragraph = doc.add_paragraph(style=list_style)
            
            # 设置缩进级别
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.5 * level)
            
            # 处理li的所有内容，包括文本和嵌套列表
            self._process_li_word_content(paragraph, li, list_style, level)

    def _process_li_word_content(self, paragraph, li_tag, parent_style, level=0):
        """处理li标签内容（Word文档生成）"""
        # 处理li标签的所有子元素
        for child in li_tag.children:
            if hasattr(child, 'name'):
                if child.name == 'strong':
                    # 处理粗体文本
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['ul', 'ol']:
                    # 处理嵌套列表 - 创建新的段落并设置缩进
                    self._process_nested_list_separate(child, parent_style, level + 1)
                else:
                    # 处理其他标签
                    paragraph.add_run(child.get_text())
            elif isinstance(child, NavigableString) and child.strip():
                # 处理纯文本
                paragraph.add_run(child.strip())

    def _process_nested_list_separate(self, list_tag, parent_style, level):
        """处理嵌套列表（单独段落）"""
        # 为嵌套列表创建新的段落
        for li in list_tag.find_all('li', recursive=False):
            # 创建新的段落
            paragraph = self.doc.add_paragraph(style=parent_style)
            
            # 设置段落缩进
            paragraph.paragraph_format.left_indent = Inches(0.5 * level)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            
            # 处理列表项内容
            self._process_li_word_content(paragraph, li, parent_style, level)
            
            # 递归处理更深层的嵌套列表
            for child in li.children:
                if hasattr(child, 'name') and child.name in ['ul', 'ol']:
                    self._process_nested_list_separate(child, parent_style, level + 1)

    def to_word(self,file_name):
        """
        将markdown文本转换成word，
        :param file_name
        :return: 返回文件名
        """
        if file_name is None or file_name =='':
            file_name = CommonUtils.generate_uuid() + ".docx"
        elif not file_name.endswith(".docx"):
            file_name += ".docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img','ul','ol']):
            if tag.name == 'h1':
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h5':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h6':
                heading = doc.add_heading(level=4)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'p':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(tag.get_text())
                run.font.size = Pt(12)
            elif tag.name == 'img':
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        fixed_width = Inches(6)

                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'ul':
                self._process_list(doc, tag, 'ListBullet', 0)
            elif tag.name == 'ol':
                self._process_list(doc, tag, 'ListNumber', 0)
            elif tag.name == 'table':
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for cell in table.rows[0].cells:
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1

                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,file_name,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param file_name:
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        if file_name is None or file_name=='':
            file_name = CommonUtils.generate_uuid()+ ".pdf";
        elif not file_name.endswith(".pdf"):
            file_name += ".pdf"
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name
