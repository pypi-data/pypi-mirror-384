from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config
import re

class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = self.format_markdown(markdown_content)
        print(self.markdown_content)
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        print(self.html_content)
        # 逐行处理HTML代码，在<ol>、<li>、<ul>标签前添加换行
        self.html_content = self._process_html_line_breaks(self.html_content)
        
        FileUtils.create_paths(Config.BASE_PATH)

    def format_markdown(self, markdown_text):
        """
        逐行处理Markdown格式，确保格式正确，特别是嵌套列表结构
        
        Args:
            markdown_text (str): 原始Markdown文本
            
        Returns:
            str: 格式化后的Markdown文本
        """
        lines = markdown_text.split('\n')
        processed_lines = []
        
        # 跟踪列表嵌套级别
        list_stack = []
        
        for i, line in enumerate(lines):
            line = line.rstrip()  # 移除行尾空白
            
            # 跳过空行
            if not line.strip():
                processed_lines.append('')
                continue
            
            # 处理标题行：去掉前面的空格，确保#号后面有空格
            if line.strip().startswith('#'):
                line = re.sub(r'^\s*(#+)\s*', r'\1 ', line)
                # 标题重置列表嵌套
                list_stack = []
            
            # 处理表格和标题之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and line.strip().startswith('#'):
                processed_lines.append('')  # 添加空行分隔
                list_stack = []
            
            # 处理表格和列表之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and re.match(r'^\s*(\d+\.|\-)', line.strip()):
                processed_lines.append('')  # 添加空行分隔
                list_stack = []
            
            # 处理列表项 - 确保正确的嵌套结构
            list_match = re.match(r'^(\s*)(\d+\.|\-)\s+', line)
            if list_match:
                indent = len(list_match.group(1))
                marker = list_match.group(2)
                
                # 确定当前缩进级别
                current_level = indent // 2  # 每2个空格一个级别
                
                # 调整列表嵌套
                while len(list_stack) > current_level:
                    list_stack.pop()
                
                # 如果当前级别大于堆栈大小，添加新的列表级别
                while len(list_stack) < current_level:
                    list_stack.append(marker)
                
                # 确保列表项格式正确
                if marker == '-':
                    # 项目符号列表
                    line = '  ' * current_level + '- ' + line[len(list_match.group(0)):].strip()
                else:
                    # 数字列表 - 保留原始序号
                    line = '  ' * current_level + marker + ' ' + line[len(list_match.group(0)):].strip()
            
            # 处理冒号后的换行 - 关键修复：确保嵌套列表有正确的换行
            if '：' in line and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # 如果下一行是列表项，在当前行后添加空行
                if re.match(r'^\s*(\d+\.|\-)', next_line):
                    line = line + '\n'
            
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)

    def to_html(self):
        return self.html_content

    def _process_html_line_breaks(self, html_content):
        """
        逐行处理HTML代码，判断每一行末尾是否包含<ol>、<li>或<ul>标签
        如果是的话，在这些标签之前添加换行，处理异常HTML结构
        
        Args:
            html_content (str): 原始HTML内容
            
        Returns:
            str: 处理后的HTML内容
        """
        lines = html_content.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.rstrip()  # 移除行尾空白
            
            # 检查行中是否包含<ol>、<li>或<ul>标签
            # 使用正则表达式查找这些标签，并确保它们不在行首
            pattern = r'(.*?)(<(ol|li|ul)[^>]*>)'
            match = re.search(pattern, line)
            
            if match and match.group(1):  # 如果标签不在行首
                # 分割行：标签前的内容和标签本身
                before_tag = match.group(1).rstrip()
                tag_content = match.group(2)
                
                # 如果标签前有内容，则在该标签前添加换行
                if before_tag:
                    processed_lines.append(before_tag)
                    processed_lines.append(tag_content)
                else:
                    processed_lines.append(line)
            else:
                processed_lines.append(line)
        
        return '\n'.join(processed_lines)

    def _process_list(self, doc, list_tag, list_style, level=0):
        """处理列表标签，支持嵌套列表"""
        for li in list_tag.find_all('li', recursive=False):
            # 创建列表项段落
            paragraph = doc.add_paragraph(style=list_style)
            
            # 设置缩进级别
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.5 * level)
            
            # 处理li的所有内容，包括文本和嵌套列表
            self._process_li_word_content(paragraph, li, list_style, level)

    def _process_li_word_content(self, paragraph, li_tag, parent_style, level=0):
        """处理li标签内容（Word文档生成）"""
        # 处理li标签的所有子元素
        for child in li_tag.children:
            if hasattr(child, 'name'):
                if child.name == 'strong':
                    # 处理粗体文本
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['ul', 'ol']:
                    # 处理嵌套列表 - 创建新的段落并设置缩进
                    self._process_nested_list_separate(child, parent_style, level + 1)
                else:
                    # 处理其他标签
                    paragraph.add_run(child.get_text())
            elif isinstance(child, NavigableString) and child.strip():
                # 处理纯文本
                paragraph.add_run(child.strip())

    def _process_nested_list_separate(self, list_tag, parent_style, level):
        """处理嵌套列表（单独段落）"""
        # 为嵌套列表创建新的段落
        for li in list_tag.find_all('li', recursive=False):
            # 创建新的段落
            paragraph = self.doc.add_paragraph(style=parent_style)
            
            # 设置段落缩进
            paragraph.paragraph_format.left_indent = Inches(0.5 * level)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            
            # 处理列表项内容
            self._process_li_word_content(paragraph, li, parent_style, level)
            
            # 递归处理更深层的嵌套列表
            for child in li.children:
                if hasattr(child, 'name') and child.name in ['ul', 'ol']:
                    self._process_nested_list_separate(child, parent_style, level + 1)

    def to_word(self,file_name):
        """
        将markdown文本转换成word，
        :param file_name
        :return: 返回文件名
        """
        if file_name is None or file_name =='':
            file_name = CommonUtils.generate_uuid() + ".docx"
        elif not file_name.endswith(".docx"):
            file_name += ".docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)
        print(self.html_content)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img','ul','ol']):
            if tag.name == 'h1':
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h5':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h6':
                heading = doc.add_heading(level=4)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'p':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(tag.get_text())
                run.font.size = Pt(12)
            elif tag.name == 'img':
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        fixed_width = Inches(6)

                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'ul':
                self._process_list(doc, tag, 'ListBullet', 0)
            elif tag.name == 'ol':
                self._process_list(doc, tag, 'ListNumber', 0)
            elif tag.name == 'table':
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for cell in table.rows[0].cells:
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1

                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,file_name,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param file_name:
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        if file_name is None or file_name=='':
            file_name = CommonUtils.generate_uuid()+ ".pdf";
        elif not file_name.endswith(".pdf"):
            file_name += ".pdf"
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name
