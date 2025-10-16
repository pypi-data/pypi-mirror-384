from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config
import re

class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = self.format_markdown(markdown_content)
        print(self.markdown_content)
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        print(self.html_content)
        # 逐行处理HTML代码，在<ol>、<li>、<ul>标签前添加换行
        self.html_content = self._process_html_line_breaks(self.html_content)
        print('='*50)
        print(self.html_content)
        FileUtils.create_paths(Config.BASE_PATH)

    def format_markdown(self, markdown_text):
        """
        简化Markdown格式处理，只处理基本格式问题
        
        Args:
            markdown_text (str): 原始Markdown文本
            
        Returns:
            str: 格式化后的Markdown文本
        """
        lines = markdown_text.split('\n')
        processed_lines = []
        
        for i, line in enumerate(lines):
            line = line.rstrip()  # 移除行尾空白
            
            # 跳过空行
            if not line.strip():
                processed_lines.append('')
                continue
            
            # 处理标题行：去掉前面的空格，确保#号后面有空格
            if line.strip().startswith('#'):
                line = re.sub(r'^\s*(#+)\s*', r'\1 ', line)
            
            # 处理表格和标题之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and line.strip().startswith('#'):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理表格和列表之间的分隔
            if i > 0 and lines[i-1].strip().startswith('|') and re.match(r'^\s*(\d+\.|\-)', line.strip()):
                processed_lines.append('')  # 添加空行分隔
            
            # 处理冒号后的换行 - 确保嵌套列表有正确的换行
            if '：' in line and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # 如果下一行是列表项，在当前行后添加空行
                if re.match(r'^\s*(\d+\.|\-)', next_line):
                    line = line + '\n'
            
            # 特殊处理：对于以数字开头且后面是粗体的行，进行转义
            # 避免被Markdown库误识别为有序列表项
            if re.match(r'^\s*\d+\.\s*\*\*', line):
                # 转义数字点号，使其不被识别为列表项
                line = re.sub(r'^(\s*)(\d+)\.(\s*\*\*)', r'\1\2\\.\3', line)
                
                # 确保转义后的行后面有适当的换行，避免HTML结构问题
                if i + 1 < len(lines) and re.match(r'^\s*\-\s*\*\*', lines[i + 1]):
                    line = line + '\n\n'  # 添加两个换行，确保正确的HTML结构
            
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)

    def to_html(self):
        return self.html_content

    def _process_html_line_breaks(self, html_content):
        """
        处理HTML代码，去除<li>、<ul>、<ol>标签内的<p>标签，并处理换行结构
        
        Args:
            html_content (str): 原始HTML内容
            
        Returns:
            str: 处理后的HTML内容
        """
        # 逐个处理<li>标签内的<p>和</p>标签
        # 使用更精确的方法，避免统一正则表达式的问题
        
        # 方法1：处理完整的<li><p>...</p></li>结构
        def process_li_tags(match):
            li_content = match.group(1)
            # 如果li_content包含<p>和</p>标签，移除它们
            if '<p>' in li_content and '</p>' in li_content:
                # 移除<p>和</p>标签，但保留内容
                li_content = re.sub(r'<p>(.*?)</p>', r'\1', li_content, flags=re.DOTALL)
            # 处理不完整的结构：只有</p>没有<p>
            elif '</p>' in li_content:
                li_content = li_content.replace('</p>', '')
            return f'<li>{li_content}</li>'
        
        # 使用更精确的正则匹配每个<li>标签
        html_content = re.sub(r'<li>(.*?)</li>', process_li_tags, html_content, flags=re.DOTALL)    
        return html_content

    def _process_list(self, doc, list_tag, list_style, level=0, start_number=1):
        """处理列表标签，支持嵌套列表"""
        # 对于有序列表，手动设置起始序号
        if list_style == 'ListNumber':
            # 为每个有序列表项手动编号
            for i, li in enumerate(list_tag.find_all('li', recursive=False), start=start_number):
                # 创建列表项段落
                paragraph = doc.add_paragraph()
                
                # 设置缩进级别
                if level > 0:
                    paragraph.paragraph_format.left_indent = Inches(0.5 * level)
                
                # 手动添加序号
                run = paragraph.add_run(f"{i}. ")
                run.bold = False
                
                # 处理li的所有内容，包括文本和嵌套列表
                self._process_li_word_content(paragraph, li, list_style, level, doc)
        else:
            # 无序列表保持原有逻辑
            for li in list_tag.find_all('li', recursive=False):
                # 创建列表项段落
                paragraph = doc.add_paragraph(style=list_style)
                
                # 设置缩进级别
                if level > 0:
                    paragraph.paragraph_format.left_indent = Inches(0.5 * level)
                
                # 处理li的所有内容，包括文本和嵌套列表
                self._process_li_word_content(paragraph, li, list_style, level, doc)

    def _process_li_word_content(self, paragraph, li_tag, parent_style, level=0, doc=None):
        """处理li标签内容（Word文档生成）"""
        # 处理li标签的所有子元素
        for child in li_tag.children:
            if hasattr(child, 'name'):
                if child.name == 'strong':
                    # 处理粗体文本
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['ul', 'ol']:
                    # 处理嵌套列表 - 创建新的段落并设置缩进
                    # 对于嵌套的有序列表，确保从1开始编号
                    nested_style = 'ListBullet' if child.name == 'ul' else 'ListNumber'
                    self._process_nested_list_separate(child, nested_style, level + 1, doc)
                else:
                    # 处理其他标签
                    paragraph.add_run(child.get_text())
            elif isinstance(child, NavigableString) and child.strip():
                # 处理纯文本
                paragraph.add_run(child.strip())

    def _process_nested_list_separate(self, list_tag, parent_style, level, doc):
        """处理嵌套列表（单独段落）"""
        # 为嵌套列表创建新的段落
        # 对于嵌套的有序列表，确保从1开始编号
        if parent_style == 'ListNumber':
            # 手动处理有序列表
            for i, li in enumerate(list_tag.find_all('li', recursive=False), start=1):
                # 创建新的段落
                paragraph = doc.add_paragraph()
                
                # 设置段落缩进
                paragraph.paragraph_format.left_indent = Inches(0.5 * level)
                paragraph.paragraph_format.first_line_indent = Inches(0)
                
                # 手动添加序号
                run = paragraph.add_run(f"{i}. ")
                run.bold = False
                
                # 处理列表项内容
                self._process_li_word_content(paragraph, li, parent_style, level, doc)
                
                # 递归处理更深层的嵌套列表
                for child in li.children:
                    if hasattr(child, 'name') and child.name in ['ul', 'ol']:
                        nested_style = 'ListBullet' if child.name == 'ul' else 'ListNumber'
                        self._process_nested_list_separate(child, nested_style, level + 1, doc)
        else:
            # 无序列表保持原有逻辑
            for li in list_tag.find_all('li', recursive=False):
                # 创建新的段落
                paragraph = doc.add_paragraph(style=parent_style)
                
                # 设置段落缩进
                paragraph.paragraph_format.left_indent = Inches(0.5 * level)
                paragraph.paragraph_format.first_line_indent = Inches(0)
                
                # 处理列表项内容
                self._process_li_word_content(paragraph, li, parent_style, level, doc)
                
                # 递归处理更深层的嵌套列表
                for child in li.children:
                    if hasattr(child, 'name') and child.name in ['ul', 'ol']:
                        nested_style = 'ListBullet' if child.name == 'ul' else 'ListNumber'
                        self._process_nested_list_separate(child, nested_style, level + 1, doc)

    def to_word(self,file_name):
        """
        将markdown文本转换成word，
        :param file_name
        :return: 返回文件名
        """
        if file_name is None or file_name =='':
            file_name = CommonUtils.generate_uuid() + ".docx"
        elif not file_name.endswith(".docx"):
            file_name += ".docx"
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img','ul','ol']):
            if tag.name == 'h1':
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h5':
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h6':
                heading = doc.add_heading(level=4)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'p':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(tag.get_text())
                run.font.size = Pt(12)
            elif tag.name == 'img':
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        fixed_width = Inches(6)

                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'ul':
                self._process_list(doc, tag, 'ListBullet', 0)
            elif tag.name == 'ol':
                self._process_list(doc, tag, 'ListNumber', 0)
            elif tag.name == 'table':
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                for cell in table.rows[0].cells:
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1

                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,file_name,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param file_name:
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        if file_name is None or file_name=='':
            file_name = CommonUtils.generate_uuid()+ ".pdf";
        elif not file_name.endswith(".pdf"):
            file_name += ".pdf"
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name
